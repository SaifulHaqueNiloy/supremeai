# backend/api/routes/chat_export.py
"""Feature S7: Export conversations to PDF, Markdown, or Word formats.

Provides endpoints to download conversation history in multiple file formats.
Uses the existing Supabase client for data access and reportlab / python-docx
when available for document generation.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime
from typing import Any

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from loguru import logger
from pydantic import BaseModel, Field

from api.deps import get_current_user_token
from database.supabase_client import db as supabase_db

router = APIRouter(
    prefix="/api/chat",
    tags=["Chat Export"],
    dependencies=[Depends(get_current_user_token)],
)

SUPPORTED_FORMATS = ["markdown", "pdf", "docx"]


class ExportRequest(BaseModel):
    """Request body for exporting a conversation."""

    conversation_id: str = Field(..., description="UUID of the conversation to export")
    format: str = Field(..., description="Export format: 'markdown', 'pdf', or 'docx'")


class ExportFormatInfo(BaseModel):
    """Information about a supported export format."""

    format: str
    label: str
    mime_type: str
    extension: str


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


async def _fetch_conversation(user_id: str, conversation_id: str) -> dict[str, Any]:
    """Fetch a single conversation owned by *user_id*, or raise 404."""
    if not supabase_db.client:
        raise HTTPException(status_code=503, detail="Database is not available.")

    try:
        resp = (
            supabase_db.client.table("conversations")
            .select("*")
            .eq("id", conversation_id)
            .eq("user_id", user_id)
            .execute()
        )
        rows = resp.data or []
        if not rows:
            raise HTTPException(
                status_code=404,
                detail=f"Conversation {conversation_id} not found.",
            )
        return rows[0]
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"Failed to fetch conversation {conversation_id}: {exc}")
        raise HTTPException(status_code=500, detail="Failed to fetch conversation.") from exc


async def _fetch_messages(conversation_id: str) -> list[dict[str, Any]]:
    """Fetch all messages for a conversation, ordered chronologically."""
    if not supabase_db.client:
        raise HTTPException(status_code=503, detail="Database is not available.")

    try:
        resp = (
            supabase_db.client.table("messages")
            .select("*")
            .eq("conversation_id", conversation_id)
            .order("created_at", desc=False)
            .execute()
        )
        return resp.data or []
    except Exception as exc:
        logger.error(f"Failed to fetch messages for conversation {conversation_id}: {exc}")
        raise HTTPException(
            status_code=500, detail="Failed to fetch conversation messages."
        ) from exc


def _format_timestamp(ts: str | None) -> str:
    """Parse an ISO timestamp and return a human-readable string."""
    if not ts:
        return "Unknown time"
    try:
        dt = datetime.fromisoformat(ts)
        return dt.strftime("%Y-%m-%d %H:%M:%S UTC")
    except (ValueError, TypeError):
        return ts


def _build_markdown(conversation: dict[str, Any], messages: list[dict[str, Any]]) -> str:
    """Render conversation + messages into clean Markdown."""
    title = conversation.get("title") or "Untitled Conversation"
    created = _format_timestamp(conversation.get("created_at"))
    updated = _format_timestamp(conversation.get("updated_at"))

    lines: list[str] = []
    lines.append(f"# {title}\n")
    lines.append(f"**Created:** {created}  ")
    lines.append(f"**Last Updated:** {updated}\n")
    lines.append("---\n")

    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        ts = _format_timestamp(msg.get("created_at"))
        content = msg.get("content", "")
        lines.append(f"### {role}\n")
        lines.append(f"*{ts}*\n")
        lines.append(f"{content}\n")
        lines.append("---\n")

    return "\n".join(lines)


def _build_pdf_bytes(conversation: dict[str, Any], messages: list[dict[str, Any]]) -> bytes:
    """Generate a PDF using reportlab. Falls back to markdown-wrapped PDF if
    reportlab is not installed."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError:
        logger.warning("reportlab not installed; falling back to markdown-as-PDF.")
        md_content = _build_markdown(conversation, messages).encode("utf-8")
        return md_content

    title = conversation.get("title") or "Untitled Conversation"
    created = _format_timestamp(conversation.get("created_at"))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ConvTitle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=4 * mm,
    )
    meta_style = ParagraphStyle(
        "ConvMeta",
        parent=styles["Normal"],
        fontSize=10,
        textColor="grey",
    )
    role_style = ParagraphStyle(
        "MsgRole",
        parent=styles["Heading3"],
        fontSize=13,
        spaceBefore=6 * mm,
        spaceAfter=1 * mm,
    )
    body_style = ParagraphStyle(
        "MsgBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
    )

    story: list[Any] = [
        Paragraph(title, title_style),
        Paragraph(f"Created: {created}", meta_style),
        Spacer(1, 4 * mm),
    ]

    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        content = msg.get("content", "")
        # Escape basic XML entities for reportlab
        safe_content = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(role, role_style))
        story.append(Paragraph(safe_content, body_style))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def _build_docx_bytes(conversation: dict[str, Any], messages: list[dict[str, Any]]) -> bytes:
    """Generate a .docx file. Uses python-docx when available, otherwise
    returns the markdown content with a .docx extension as a graceful fallback."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.warning(
            "python-docx not installed; falling back to markdown content with .docx extension."
        )
        return _build_markdown(conversation, messages).encode("utf-8")

    title = conversation.get("title") or "Untitled Conversation"
    created = _format_timestamp(conversation.get("created_at"))

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    if heading.runs:
        heading.runs[0].font.size = Pt(18)

    # Meta
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"Created: {created}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph("\u2014" * 40)

    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        content = msg.get("content", "")
        doc.add_heading(role, level=3)
        doc.add_paragraph(content)
        doc.add_paragraph("\u2014" * 20)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Format metadata mapping
# ---------------------------------------------------------------------------

_FORMAT_MAP: dict[str, dict[str, str]] = {
    "markdown": {
        "mime": "text/markdown; charset=utf-8",
        "ext": "md",
        "label": "Markdown",
    },
    "pdf": {
        "mime": "application/pdf",
        "ext": "pdf",
        "label": "PDF",
    },
    "docx": {
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "ext": "docx",
        "label": "Word Document",
    },
}


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.get(
    "/export/formats",
    response_model=list[ExportFormatInfo],
    summary="List available export formats",
)
async def list_export_formats() -> list[ExportFormatInfo]:
    """Return metadata about every supported export format."""
    return [
        ExportFormatInfo(
            format=fmt,
            label=meta["label"],
            mime_type=meta["mime"],
            extension=meta["ext"],
        )
        for fmt, meta in _FORMAT_MAP.items()
    ]


@router.post(
    "/export",
    summary="Export a conversation to a downloadable file",
)
async def export_conversation(
    payload: ExportRequest,
    user: dict = Depends(get_current_user_token),
) -> StreamingResponse:
    """Export a conversation in the requested format and return a file download.

    Supported formats: ``markdown``, ``pdf``, ``docx``.
    """
    user_id = user.get("sub")
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid token.")

    fmt = payload.format.lower()
    if fmt not in _FORMAT_MAP:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format '{payload.format}'. Supported: {SUPPORTED_FORMATS}",
        )

    conversation = await _fetch_conversation(user_id, payload.conversation_id)
    messages = await _fetch_messages(payload.conversation_id)

    if not messages:
        raise HTTPException(
            status_code=404,
            detail="No messages found in this conversation.",
        )

    # Build file content
    safe_title = (conversation.get("title") or "conversation").replace("/", "-").replace("\\", "-")

    if fmt == "markdown":
        content = _build_markdown(conversation, messages).encode("utf-8")
    elif fmt == "pdf":
        content = _build_pdf_bytes(conversation, messages)
    else:
        content = _build_docx_bytes(conversation, messages)

    meta = _FORMAT_MAP[fmt]
    filename = f"{safe_title}.{meta['ext']}"

    return StreamingResponse(
        io.BytesIO(content),
        media_type=meta["mime"],
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
    )
