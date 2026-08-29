"""Feature S5: Slash Commands Backend.

Provides slash command definitions and execution endpoints.
Built-in commands include /research, /summarize, /image, /code,
/translate, /think, /export, /clear, and /help.
"""

from __future__ import annotations

import asyncio
import io
import uuid
from typing import Any

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field

from api.deps import get_current_user_token
from core.logging_config import logger

router = APIRouter(
    prefix="/api/commands",
    tags=["Slash Commands"],
    dependencies=[Depends(get_current_user_token)],
)


# ---------- Command Definitions ----------


class CommandParameter(BaseModel):
    name: str = Field(..., description="Parameter name")
    type: str = Field(default="string", description="Parameter type (string, number, boolean)")
    required: bool = Field(default=False, description="Whether the parameter is required")
    description: str = Field(default="", description="Parameter description")
    default: str | None = Field(default=None, description="Default value if not provided")


class CommandDefinition(BaseModel):
    id: str
    name: str
    description: str
    icon: str
    category: str
    parameters: list[CommandParameter] | None = None


COMMAND_REGISTRY: list[CommandDefinition] = [
    CommandDefinition(
        id="research",
        name="/research",
        description="Deep research on a topic using knowledge base and web search",
        icon="🔍",
        category="Research",
        parameters=[
            CommandParameter(
                name="query",
                type="string",
                required=True,
                description="The research query",
            ),
            CommandParameter(
                name="depth",
                type="string",
                required=False,
                description="Research depth: quick, standard, deep",
                default="standard",
            ),
        ],
    ),
    CommandDefinition(
        id="summarize",
        name="/summarize",
        description="Summarize the current conversation or a specific topic",
        icon="📝",
        category="Conversation",
        parameters=[
            CommandParameter(
                name="conversation_id",
                type="string",
                required=False,
                description="Conversation ID to summarize (defaults to current)",
            ),
        ],
    ),
    CommandDefinition(
        id="image",
        name="/image",
        description="Generate an image from a text prompt",
        icon="🖼️",
        category="Creative",
        parameters=[
            CommandParameter(
                name="prompt",
                type="string",
                required=True,
                description="Image generation prompt",
            ),
            CommandParameter(
                name="model",
                type="string",
                required=False,
                description="Model to use for generation",
                default="stabilityai/stable-diffusion-xl-base-1.0",
            ),
        ],
    ),
    CommandDefinition(
        id="code",
        name="/code",
        description="Switch to code mode with a specific framework or language",
        icon="💻",
        category="Mode",
        parameters=[
            CommandParameter(
                name="framework",
                type="string",
                required=False,
                description="Target framework (react, python, nextjs, etc.)",
                default="python",
            ),
        ],
    ),
    CommandDefinition(
        id="translate",
        name="/translate",
        description="Translate text to a target language",
        icon="🌐",
        category="Utility",
        parameters=[
            CommandParameter(
                name="text",
                type="string",
                required=True,
                description="Text to translate",
            ),
            CommandParameter(
                name="target_language",
                type="string",
                required=True,
                description="Target language (e.g., Spanish, French, Japanese)",
            ),
        ],
    ),
    CommandDefinition(
        id="think",
        name="/think",
        description="Enable reasoning mode for the next message",
        icon="🧠",
        category="Mode",
        parameters=[
            CommandParameter(
                name="mode",
                type="string",
                required=False,
                description="Reasoning mode: tree_of_thought, debate, quick",
                default="tree_of_thought",
            ),
        ],
    ),
    CommandDefinition(
        id="export",
        name="/export",
        description="Export the current conversation as PDF, Markdown, or Word",
        icon="📤",
        category="Export",
        parameters=[
            CommandParameter(
                name="conversation_id",
                type="string",
                required=False,
                description="Conversation ID to export",
            ),
            CommandParameter(
                name="format",
                type="string",
                required=False,
                description="Export format: markdown, pdf, docx",
                default="markdown",
            ),
        ],
    ),
    CommandDefinition(
        id="clear",
        name="/clear",
        description="Clear the current conversation history",
        icon="🗑️",
        category="Conversation",
        parameters=[
            CommandParameter(
                name="conversation_id",
                type="string",
                required=False,
                description="Conversation ID to clear (defaults to current)",
            ),
        ],
    ),
    CommandDefinition(
        id="help",
        name="/help",
        description="Show all available slash commands",
        icon="❓",
        category="Meta",
        parameters=[],
    ),
]


# ---------- Request / Response Schemas ----------


class CommandExecuteRequest(BaseModel):
    command: str = Field(..., description="Slash command name, e.g. /research")
    args: dict[str, Any] = Field(default_factory=dict, description="Command arguments")
    conversation_id: str | None = Field(default=None, description="Optional conversation context")


class CommandExecuteResponse(BaseModel):
    command: str
    status: str
    result: Any
    message: str


# ---------- Command Handlers ----------


async def _handle_research(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Execute deep research using the knowledge base and LLM."""
    query = args.get("query", "")
    depth = args.get("depth", "standard")

    if not query:
        return {"error": "Query is required for /research"}

    # Search knowledge base
    kb_results = []
    try:
        from tools.knowledge.knowledge_base_indexer import KnowledgeBaseIndexer

        indexer = KnowledgeBaseIndexer()
        n_results = {"quick": 3, "standard": 5, "deep": 10}.get(depth, 5)
        kb_results = indexer.search_knowledge(query, n_results=n_results)
    except Exception as e:
        logger.warning(f"Knowledge base search failed: {e}")

    # Use LLM to synthesize research
    try:
        from core.llm.llm_gateway import llm_gateway

        kb_context = ""
        if kb_results:
            kb_context = "\n".join(
                f"- [{r.get('metadata', {}).get('type', 'unknown')}] {r.get('text', '')[:300]}"
                for r in kb_results[:5]
            )

        prompt = (
            f"Conduct a thorough research analysis on: {query}\n\n"
            f"Relevant knowledge base findings:\n{kb_context or 'No prior knowledge found.'}\n\n"
            f"Provide a comprehensive research summary with key findings, "
            f"actionable insights, and recommended next steps."
        )

        response = await llm_gateway.acompletion(prompt=prompt, task_type="research", stream=False)
        summary = response.get("text", "") if isinstance(response, dict) else str(response)
    except Exception as e:
        logger.warning(f"LLM research synthesis failed: {e}")
        summary = f"Research query: {query}. Knowledge base returned {len(kb_results)} results. LLM synthesis unavailable."

    return {
        "query": query,
        "depth": depth,
        "kb_results_count": len(kb_results),
        "summary": summary,
    }


async def _handle_summarize(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Summarize a conversation."""
    from database.supabase_client import SupabaseDB

    conversation_id = args.get("conversation_id")
    db = SupabaseDB()

    # If no conversation_id provided, get the most recent
    if not conversation_id:
        convs = (
            await db.client.table("conversations")
            .select("id")
            .eq("user_id", user_id)
            .order("updated_at", desc=True)
            .limit(1)
            .execute()
        )
        if not convs.data:
            return {"error": "No conversations found"}
        conversation_id = convs.data[0]["id"]

    # Fetch messages
    messages_resp = (
        await db.client.table("messages")
        .select("role, content")
        .eq("conversation_id", conversation_id)
        .order("created_at", desc=False)
        .execute()
    )

    if not messages_resp.data:
        return {"conversation_id": conversation_id, "summary": "This conversation has no messages."}

    # Build conversation text
    lines = []
    for msg in messages_resp.data:
        role = msg["role"].capitalize()
        content = msg["content"][:500]
        lines.append(f"{role}: {content}")

    conversation_text = "\n".join(lines)

    # Summarize via LLM
    try:
        from core.llm.llm_gateway import llm_gateway

        prompt = (
            f"Summarize the following conversation concisely, capturing the key points, "
            f"decisions made, and any action items:\n\n{conversation_text}"
        )
        response = await llm_gateway.acompletion(
            prompt=prompt, task_type="summarization", stream=False
        )
        summary = response.get("text", "") if isinstance(response, dict) else str(response)
    except Exception as e:
        logger.warning(f"Summarization LLM call failed: {e}")
        # Fallback: first 200 chars
        summary = conversation_text[:200] + "..."

    return {
        "conversation_id": conversation_id,
        "message_count": len(messages_resp.data),
        "summary": summary,
    }


async def _handle_image(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Generate an image using the HF image generator."""
    prompt = args.get("prompt", "")
    model = args.get("model")

    if not prompt:
        return {"error": "Prompt is required for /image"}

    try:
        from tools.media.image_generator import HFImageGenerator

        generator = HFImageGenerator()
        output_path = f"data/generated_{uuid.uuid4().hex[:8]}.png"
        result = await generator.generate_image(prompt=prompt, model=model, output_path=output_path)
        return result
    except Exception as e:
        logger.error(f"Image generation failed: {e}")
        return {"success": False, "error": str(e)}


async def _handle_code(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Switch to code mode with a specific framework."""
    framework = args.get("framework", "python")
    return {
        "mode": "code",
        "framework": framework,
        "message": f"Code mode activated with {framework} framework. Your next message will be treated as code instructions.",
    }


async def _handle_translate(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Translate text to a target language."""
    text = args.get("text", "")
    target = args.get("target_language", "")

    if not text:
        return {"error": "Text is required for /translate"}
    if not target:
        return {"error": "Target language is required for /translate"}

    try:
        from core.llm.llm_gateway import llm_gateway

        prompt = (
            f"Translate the following text to {target}. "
            f"Only provide the translated text, nothing else:\n\n{text}"
        )
        response = await llm_gateway.acompletion(
            prompt=prompt, task_type="translation", stream=False
        )
        translated = response.get("text", "") if isinstance(response, dict) else str(response)
    except Exception as e:
        logger.error(f"Translation failed: {e}")
        return {"error": f"Translation failed: {e}"}

    return {
        "original": text,
        "translated": translated,
        "target_language": target,
    }


async def _handle_think(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Enable reasoning mode."""
    mode = args.get("mode", "tree_of_thought")
    valid_modes = ["tree_of_thought", "debate", "quick"]
    if mode not in valid_modes:
        mode = "tree_of_thought"
    return {
        "mode": "reasoning",
        "reasoning_engine": mode,
        "message": f"Reasoning mode enabled using '{mode}'. Your next message will be processed with enhanced reasoning.",
    }


async def _handle_export(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Export a conversation."""
    from database.supabase_client import SupabaseDB

    conversation_id = args.get("conversation_id")
    fmt = args.get("format", "markdown")
    db = SupabaseDB()

    if not conversation_id:
        convs = (
            await db.client.table("conversations")
            .select("id, title")
            .eq("user_id", user_id)
            .order("updated_at", desc=True)
            .limit(1)
            .execute()
        )
        if not convs.data:
            return {"error": "No conversations found to export"}
        conversation_id = convs.data[0]["id"]
        title = convs.data[0].get("title", "Export")
    else:
        conv = (
            await db.client.table("conversations")
            .select("title")
            .eq("id", conversation_id)
            .execute()
        )
        title = conv.data[0]["title"] if conv.data else "Export"

    # Fetch messages
    msgs = (
        db.client.table("messages")
        .select("role, content, created_at")
        .eq("conversation_id", conversation_id)
        .order("created_at", desc=False)
        .execute()
    )

    if not msgs.data:
        return {"error": "No messages to export"}

    if fmt == "markdown":
        lines = [f"# {title}\n"]
        for m in msgs.data:
            role = m["role"].capitalize()
            lines.append(f"### {role}\n{m['content']}\n")
        content = "\n".join(lines)
        return {"format": "markdown", "content": content, "message_count": len(msgs.data)}

    if fmt == "pdf":
        # Generate simple PDF using reportlab if available
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = [Paragraph(f"<h1>{title}</h1>", styles["Heading1"])]
            for m in msgs.data:
                role = m["role"].capitalize()
                story.append(Paragraph(f"<b>{role}:</b>", styles["Normal"]))
                story.append(Paragraph(m["content"].replace("\n", "<br/>"), styles["Normal"]))
                story.append(Paragraph("", styles["Normal"]))

            doc.build(story)
            buffer.seek(0)
            return {
                "format": "pdf",
                "content_base64": buffer.read().hex(),
                "message_count": len(msgs.data),
            }
        except ImportError:
            return {"error": "PDF export requires reportlab. Install with: pip install reportlab"}

    if fmt == "docx":
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(title, level=1)
            for m in msgs.data:
                role = m["role"].capitalize()
                doc.add_heading(role, level=2)
                doc.add_paragraph(m["content"])

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return {
                "format": "docx",
                "content_base64": buffer.read().hex(),
                "message_count": len(msgs.data),
            }
        except ImportError:
            return {
                "error": "Word export requires python-docx. Install with: pip install python-docx"
            }

    return {"error": f"Unsupported export format: {fmt}. Use markdown, pdf, or docx."}


async def _handle_clear(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Clear a conversation's messages."""
    from database.supabase_client import SupabaseDB

    conversation_id = args.get("conversation_id")
    db = SupabaseDB()

    if not conversation_id:
        convs = (
            db.client.table("conversations")
            .select("id")
            .eq("user_id", user_id)
            .order("updated_at", desc=True)
            .limit(1)
            .execute()
        )
        if not convs.data:
            return {"error": "No conversations found to clear"}
        conversation_id = convs.data[0]["id"]

    try:
        await db.client.table("messages").delete().eq("conversation_id", conversation_id).execute()
        return {
            "status": "cleared",
            "conversation_id": conversation_id,
            "message": "Conversation cleared successfully.",
        }
    except Exception as e:
        logger.error(f"Failed to clear conversation: {e}")
        return {"error": f"Failed to clear conversation: {e}"}


def _handle_help(args: dict[str, Any], user_id: str) -> dict[str, Any]:
    """Return all available commands."""
    commands = []
    for cmd in COMMAND_REGISTRY:
        entry = {
            "name": cmd.name,
            "description": cmd.description,
            "icon": cmd.icon,
            "category": cmd.category,
        }
        if cmd.parameters:
            entry["parameters"] = [
                {
                    "name": p.name,
                    "type": p.type,
                    "required": p.required,
                    "description": p.description,
                }
                for p in cmd.parameters
            ]
        commands.append(entry)
    return {"commands": commands, "total": len(commands)}


# Command name -> handler mapping
_HANDLERS: dict[str, Any] = {
    "research": _handle_research,
    "summarize": _handle_summarize,
    "image": _handle_image,
    "code": _handle_code,
    "translate": _handle_translate,
    "think": _handle_think,
    "export": _handle_export,
    "clear": _handle_clear,
    "help": _handle_help,
}


# ---------- Routes ----------


@router.get("/", response_model=list[CommandDefinition])
async def list_commands(
    user: dict = Depends(get_current_user_token),
):
    """Return all available slash commands."""
    return COMMAND_REGISTRY


@router.post("/execute", response_model=CommandExecuteResponse)
async def execute_command(
    payload: CommandExecuteRequest,
    user: dict = Depends(get_current_user_token),
):
    """Execute a slash command with the given arguments.

    Dispatches to the appropriate handler based on the command name.
    """
    user_id = user.get("sub")
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid token")

    # Normalize command name (strip leading slash)
    command_name = payload.command.lstrip("/").strip().lower()

    handler = _HANDLERS.get(command_name)
    if handler is None:
        available = ", ".join(f"/{k}" for k in _HANDLERS)
        raise HTTPException(
            status_code=400,
            detail=f"Unknown command: /{command_name}. Available commands: {available}",
        )

    logger.info(f"Executing slash command: /{command_name} for user {user_id}")

    try:
        if asyncio.iscoroutinefunction(handler):
            result = await handler(payload.args, user_id)
        else:
            result = handler(payload.args, user_id)
    except Exception as e:
        logger.error(f"Command /{command_name} execution failed: {e}")
        raise HTTPException(status_code=500, detail=f"Command execution failed: {e}") from e

    # Check if handler returned an error
    if isinstance(result, dict) and "error" in result:
        return CommandExecuteResponse(
            command=payload.command,
            status="error",
            result=result,
            message=result["error"],
        )

    return CommandExecuteResponse(
        command=payload.command,
        status="success",
        result=result,
        message=f"Command /{command_name} executed successfully.",
    )
